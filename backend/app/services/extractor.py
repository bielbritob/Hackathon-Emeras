import io
import os
import pandas as pd
import fitz
from docx import Document
import pytesseract
from PIL import Image
from fastapi import UploadFile

class DocumentExtractorService:
    """
    Comentário: Serviço responsável por processar múltiplos formatos de arquivo,
    incluindo extração de texto em imagens via OCR (Tesseract).
    """

    def __init__(self, file: UploadFile):
        # Comentário: Atributos privados protegidos por name mangling
        self.__file = file
        self.__file_name = file.filename or "desconhecido"
        self.__file_extension = os.path.splitext(self.__file_name)[1].lower()
        self.__extracted_text = ""

    @property
    def extracted_text(self) -> str:
        return self.__extracted_text

    @property
    def file_name(self) -> str:
        return self.__file_name

    async def process_document(self) -> str:
        """
        Comentário: Analisa a extensão do arquivo e direciona para o motor adequado.
        """
        content = await self.__file.read()

        if self.__file_extension == ".pdf":
            self.__extracted_text = self.__extract_from_pdf(content)
        elif self.__file_extension == ".docx":
            self.__extracted_text = self.__extract_from_docx(content)
        elif self.__file_extension in [".xlsx", ".xls"]:
            self.__extracted_text = self.__extract_from_excel(content)
        elif self.__file_extension in [".csv", ".txt"]:
            self.__extracted_text = self.__extract_from_text(content)
        elif self.__file_extension in [".png", ".jpg", ".jpeg"]:
            self.__extracted_text = self.__extract_from_image(content)
        else:
            raise ValueError(f"Extensão de arquivo '{self.__file_extension}' não é suportada.")

        return self.extracted_text

    def __extract_from_pdf(self, file_content: bytes) -> str:
        """
        Comentário: Extração de texto de arquivos PDF utilizando PyMuPDF (fitz),
        que é tolerante a falhas de EOF e muito mais rápido.
        """
        # Abre o documento PDF diretamente do fluxo de bytes em memória
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        pages_text = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text()
            if text:
                pages_text.append(text)
                
        pdf_document.close()
        return "\n".join(pages_text).strip()

    def __extract_from_docx(self, file_content: bytes) -> str:
        """
        Comentário: Extração de parágrafos e tabelas de arquivos Word (.docx).
        """
        document = Document(io.BytesIO(file_content))
        extracted_content = [p.text for p in document.paragraphs if p.text.strip()]

        # Comentário: Percorre e extrai células de tabelas dentro do documento DOCX
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    extracted_content.append(row_text)

        return "\n".join(extracted_content).strip()

    def __extract_from_excel(self, file_content: bytes) -> str:
        """
        Comentário: Extração de dados de planilhas do Excel (.xlsx, .xls) aba por aba.
        """
        excel_file = pd.ExcelFile(io.BytesIO(file_content))
        sheet_contents = []
        for sheet_name in excel_file.sheet_names:
            dataframe = pd.read_excel(excel_file, sheet_name=sheet_name)
            sheet_contents.append(f"--- Aba: {sheet_name} ---")
            sheet_contents.append(dataframe.to_string(index=False))
        return "\n".join(sheet_contents).strip()

    def __extract_from_text(self, file_content: bytes) -> str:
        """
        Comentário: Extração de arquivos de texto simples e CSV com tratamento de encoding.
        """
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            return file_content.decode("latin-1", errors="ignore")

    def __extract_from_image(self, file_content: bytes) -> str:
        """
        Comentário: Extrai texto de ficheiros de imagem (PNG, JPG, JPEG) usando OCR.
        O parâmetro lang='por' garante o reconhecimento correto de acentos (ç, ã, á).
        """
        image = Image.open(io.BytesIO(file_content))
        
        # Comentário: Executa o OCR focado no idioma Português
        extracted_content = pytesseract.image_to_string(image, lang='por')
        
        return extracted_content.strip()